"""
报告生成器
Report Generator
"""

import os
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime
import json
from pathlib import Path

# 尝试导入不同的报告生成库
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

logger = logging.getLogger(__name__)

class ReportGenerator:
    """报告生成器"""

    def __init__(self, template_dir: str = None):
        if template_dir is None:
            template_dir = Path(__file__).parent.parent.parent / 'templates'
        self.template_dir = Path(template_dir)
        self.template_dir.mkdir(exist_ok=True)
        self.field_mappings = {}
        self.load_default_mappings()

    def load_default_mappings(self):
        """加载默认字段映射"""
        self.field_mappings = {
            'vehicle_emission_report': {
                'template_file': 'vehicle_emission_template.docx',
                'fields': {
                    '[VIN码]': 'vin',
                    '[品牌]': 'make',
                    '[车型]': 'model',
                    '[发动机型号]': 'engine_code',
                    '[排量]': 'displacement',
                    '[排放标准]': 'emission_standard',
                    '[CO2排放]': 'co2_emission',
                    '[油耗]': 'fuel_consumption',
                    '[测试日期]': 'test_date',
                    '[检测员]': 'inspector'
                }
            },
            'vehicle_basic_info': {
                'template_file': 'vehicle_basic_template.docx',
                'fields': {
                    '[VIN码]': 'vin',
                    '[品牌]': 'make',
                    '[车型]': 'model',
                    '[年份]': 'year',
                    '[生产日期]': 'production_date',
                    '[发动机型号]': 'engine_code',
                    '[变速箱型号]': 'transmission_code'
                }
            }
        }

    def generate_report(self,
                       data: Dict[str, Any],
                       template_name: str,
                       output_path: str,
                       output_format: str = 'pdf') -> bool:
        """
        生成报告

        Args:
            data: 报告数据
            template_name: 模板名称
            output_path: 输出路径
            output_format: 输出格式 (pdf, docx, excel)

        Returns:
            是否生成成功
        """
        try:
            logger.info(f"开始生成报告: {template_name} -> {output_path}")

            # 验证模板
            template_config = self.field_mappings.get(template_name)
            if not template_config:
                logger.error(f"未找到模板配置: {template_name}")
                return False

            # 准备数据
            prepared_data = self._prepare_data(data, template_config)

            # 根据格式生成报告
            if output_format.lower() == 'pdf':
                success = self._generate_pdf_report(prepared_data, template_config, output_path)
            elif output_format.lower() == 'docx':
                success = self._generate_docx_report(prepared_data, template_config, output_path)
            elif output_format.lower() == 'excel':
                success = self._generate_excel_report(prepared_data, template_config, output_path)
            else:
                logger.error(f"不支持的输出格式: {output_format}")
                return False

            if success:
                logger.info(f"报告生成成功: {output_path}")
                return True
            else:
                logger.error("报告生成失败")
                return False

        except Exception as e:
            logger.error(f"生成报告时出错: {e}")
            return False

    def _prepare_data(self, data: Dict[str, Any], template_config: Dict[str, Any]) -> Dict[str, Any]:
        """准备报告数据"""
        prepared = {}
        fields = template_config['fields']

        for placeholder, field_name in fields.items():
            # 支持嵌套字段访问，如 "vehicle.vin"
            value = self._get_nested_value(data, field_name)

            # 格式化值
            formatted_value = self._format_value(value)
            prepared[placeholder] = formatted_value

        # 添加元数据
        prepared['[生成日期]'] = datetime.now().strftime('%Y年%m月%d日')
        prepared['[生成时间]'] = datetime.now().strftime('%H:%M:%S')

        return prepared

    def _get_nested_value(self, data: Dict[str, Any], field_path: str) -> Any:
        """获取嵌套字段值"""
        keys = field_path.split('.')
        value = data

        for key in keys:
            if isinstance(value, dict) and key in value:
                value = value[key]
            else:
                return None

        return value

    def _format_value(self, value: Any) -> str:
        """格式化值"""
        if value is None:
            return ""
        elif isinstance(value, datetime):
            return value.strftime('%Y年%m月%d日')
        elif isinstance(value, (int, float)):
            return str(value)
        else:
            return str(value)

    def _generate_pdf_report(self, data: Dict[str, Any], template_config: Dict[str, Any], output_path: str) -> bool:
        """生成PDF报告"""
        if not PDF_AVAILABLE:
            logger.error("PDF生成库未安装")
            return False

        try:
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []

            # 标题
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1  # 居中
            )
            story.append(Paragraph("车辆检测报告", title_style))
            story.append(Spacer(1, 20))

            # 信息表格
            table_data = []
            for placeholder, value in data.items():
                if placeholder.startswith('[') and placeholder.endswith(']'):
                    field_name = placeholder[1:-1]  # 去掉方括号
                    table_data.append([field_name, value])

            if table_data:
                table = Table(table_data, colWidths=[3*inch, 4*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)

            doc.build(story)
            return True

        except Exception as e:
            logger.error(f"生成PDF报告失败: {e}")
            return False

    def _generate_docx_report(self, data: Dict[str, Any], template_config: Dict[str, Any], output_path: str) -> bool:
        """生成Word报告"""
        if not DOCX_AVAILABLE:
            logger.error("Word生成库未安装")
            return False

        try:
            template_file = self.template_dir / template_config['template_file']

            if template_file.exists():
                # 使用模板
                doc = Document(str(template_file))
                self._replace_placeholders(doc, data)
            else:
                # 创建新文档
                doc = Document()
                self._create_docx_content(doc, data)

            doc.save(output_path)
            return True

        except Exception as e:
            logger.error(f"生成Word报告失败: {e}")
            return False

    def _replace_placeholders(self, doc, data: Dict[str, Any]):
        """替换文档中的占位符"""
        for paragraph in doc.paragraphs:
            for placeholder, value in data.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

        # 处理表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in data.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

    def _create_docx_content(self, doc, data: Dict[str, Any]):
        """创建Word文档内容"""
        # 添加标题
        doc.add_heading('车辆检测报告', 0)

        # 添加信息表格
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        for placeholder, value in data.items():
            if placeholder.startswith('[') and placeholder.endswith(']'):
                field_name = placeholder[1:-1]
                row_cells = table.add_row().cells
                row_cells[0].text = field_name
                row_cells[1].text = str(value)

    def _generate_excel_report(self, data: Dict[str, Any], template_config: Dict[str, Any], output_path: str) -> bool:
        """生成Excel报告"""
        if not EXCEL_AVAILABLE:
            logger.error("Excel生成库未安装")
            return False

        try:
            # 准备数据
            excel_data = []
            for placeholder, value in data.items():
                if placeholder.startswith('[') and placeholder.endswith(']'):
                    field_name = placeholder[1:-1]
                    excel_data.append([field_name, value])

            # 创建DataFrame并保存
            df = pd.DataFrame(excel_data, columns=['项目', '值'])
            df.to_excel(output_path, index=False, sheet_name='检测报告')
            return True

        except Exception as e:
            logger.error(f"生成Excel报告失败: {e}")
            return False

    def create_template(self, template_name: str, template_type: str, fields: Dict[str, str]) -> bool:
        """创建新模板"""
        try:
            template_config = {
                'template_file': f'{template_name}_template.{template_type}',
                'fields': fields
            }

            self.field_mappings[template_name] = template_config

            # 如果是Word模板，创建一个基础模板文件
            if template_type == 'docx' and DOCX_AVAILABLE:
                template_path = self.template_dir / template_config['template_file']
                doc = Document()
                doc.add_heading(f'{template_name}模板', 0)

                for placeholder in fields.keys():
                    doc.add_paragraph(placeholder)

                doc.save(str(template_path))

            logger.info(f"模板创建成功: {template_name}")
            return True

        except Exception as e:
            logger.error(f"创建模板失败: {e}")
            return False

    def get_template_list(self) -> List[str]:
        """获取模板列表"""
        return list(self.field_mappings.keys())

    def validate_template(self, template_name: str) -> Dict[str, Any]:
        """验证模板"""
        result = {
            'is_valid': True,
            'errors': [],
            'warnings': []
        }

        try:
            template_config = self.field_mappings.get(template_name)
            if not template_config:
                result['is_valid'] = False
                result['errors'].append(f"模板不存在: {template_name}")
                return result

            # 检查模板文件
            template_file = self.template_dir / template_config['template_file']
            if not template_file.exists():
                result['warnings'].append(f"模板文件不存在: {template_file}")

            # 检查字段映射
            if not template_config.get('fields'):
                result['warnings'].append("模板没有字段映射")

        except Exception as e:
            result['is_valid'] = False
            result['errors'].append(f"验证过程出错: {e}")

        return result

    def batch_generate_reports(self,
                              data_list: List[Dict[str, Any]],
                              template_name: str,
                              output_dir: str,
                              output_format: str = 'pdf') -> Dict[str, Any]:
        """批量生成报告"""
        results = {
            'success_count': 0,
            'failed_count': 0,
            'failed_files': [],
            'output_files': []
        }

        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)

        for i, data in enumerate(data_list):
            try:
                # 生成文件名
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"report_{i+1}_{timestamp}.{output_format}"
                file_path = output_path / filename

                # 生成报告
                success = self.generate_report(data, template_name, str(file_path), output_format)

                if success:
                    results['success_count'] += 1
                    results['output_files'].append(str(file_path))
                else:
                    results['failed_count'] += 1
                    results['failed_files'].append(f"数据项 {i+1}")

            except Exception as e:
                logger.error(f"批量生成报告时出错 (数据项 {i+1}): {e}")
                results['failed_count'] += 1
                results['failed_files'].append(f"数据项 {i+1}: {str(e)}")

        return results